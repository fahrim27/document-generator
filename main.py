from flask import Flask, render_template, request, send_file, redirect, url_for, session, jsonify
import uvicorn
import os
import json
import re
import threading
import time
from flask_session import Session
from werkzeug.utils import secure_filename
from flask_cors import CORS
import image
from docx.shared import Cm, Mm
from docxtpl import DocxTemplate, InlineImage
from time import gmtime, strftime
import urllib.request
import string, random, requests, io, shutil
from docx2pdf import convert
from spire.doc import *
from spire.doc.common import *
import fitz

ALLOWED_EXTENSIONS = set(['xls', 'csv', 'png', 'jpeg', 'jpg'])
UPLOAD_FOLDER = os.path.abspath(os.path.join(os.path.dirname(__file__), 'Downloads'))
app = Flask(__name__)
app.secret_key = 'fm27'
app.config["SESSION_PERMANENT"] = False
app.config["SESSION_TYPE"] = "filesystem"
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 500 * 1000 * 1000  # 500 MB
app.config['CORS_HEADER'] = 'application/json'

def download_and_redirect1(file_path):
    time.sleep(30)
    os.remove(file_path)
    os.remove(fr"uploads/{file_path.split('filled_')[1]}")
    return redirect('/')

def download_and_redirect2(file_path):
    time.sleep(10)
    return redirect('/')

def download_and_redirect_api(file_path, folder):
    time.sleep(15)
    os.remove(file_path)
    shutil.rmtree(folder)
    return 'success'

def download_and_redirect_pdf(pdf1, pdf2):
    time.sleep(10)
    os.remove(pdf1)
    os.remove(pdf2)
    return 'success'

def extract_words(text):
    pattern = r'\{\{(.*?)\}\}'
    matches = re.findall(pattern, text)
    return matches 

def generate_html_form(words):
    html_content = '''
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Document Auto-Fill</title>
        <style>
            body {
                background-color: #222;
                color: #fff;
                font-family: Arial, sans-serif;
            }
            .form-container {
                max-width: 800px;
                margin: 50px auto;
                padding: 20px;
                border-radius: 10px;
                background-color: #333;
                box-shadow: 0 0 10px rgba(0, 0, 0, 0.5);
            }
            .form-group {
                margin-bottom: 20px;
            }
            label {
                display: block;
                margin-bottom: 5px;
            }
            input[type="text"] {
                width: 90%;
                padding: 10px;
                border-radius: 5px;
                border: 1px solid #666;
                background-color: #444;
                color: #fff;
            }
            .textarea {
                width: 90%;
                padding: 10px;
                border-radius: 5px;
                border: 1px solid #666;
                background-color: #444;
                color: #fff;
                height:80px;
            }
            input[type="date"] {
                width: 90%;
                padding: 10px;
                border-radius: 5px;
                border: 1px solid #666;
                background-color: #444;
                color: #fff;
            }
            input[type="file"] {
                width: 90%;
                padding: 10px;
                border-radius: 5px;
                border: 1px solid #666;
                background-color: #444;
                color: #fff;
            }
            input[type="number"] {
                width: 90%;
                padding: 10px;
                border-radius: 5px;
                border: 1px solid #666;
                background-color: #444;
                color: #fff;
            }
            input[type="email"] {
                width: 90%;
                padding: 10px;
                border-radius: 5px;
                border: 1px solid #666;
                background-color: #444;
                color: #fff;
            }
            input[type="password"] {
                width: 90%;
                padding: 10px;
                border-radius: 5px;
                border: 1px solid #666;
                background-color: #444;
                color: #fff;
            }
            input[type="submit"] {
                padding: 10px 20px;
                border-radius: 5px;
                border: none;
                background-color: #007bff;
                color: #fff;
                cursor: pointer;
            }
            input[type="submit"]:hover {
                background-color: #0056b3;
            }
        </style>
    </head>
    <body>
        <div class="form-container">
            <form action="/submit" method="post" onkeydown="return event.key != 'Enter';" enctype="multipart/form-data">
    '''
    for word in sorted(words):

        if 'text' in word and 'area' not in word:
            new_word = word.replace('_', ' ')
            new_word = new_word.replace('text', '')

            inputtype = 'text'
        elif 'textarea' in word:
            new_word = word.replace('_', ' ')
            new_word = new_word.replace('textarea', '')

            inputtype = 'textarea'
        elif 'number' in word:
            new_word = word.replace('_', ' ')
            new_word = new_word.replace('number', '')

            inputtype = 'number'
        elif 'email' in word:
            new_word = word.replace('_', ' ')
            new_word = new_word.replace('email', '')

            inputtype = 'email'
        elif 'date' in word:
            new_word = word.replace('_', ' ')
            new_word = new_word.replace('date', '')

            inputtype = 'date'
        elif 'file' in word:
            new_word = word.replace('_', ' ')
            new_word = new_word.replace('file', '')
            new_word = re.sub('[0-9]', '', new_word)

            inputtype = 'file'
        else:
            new_word = word
            inputtype = 'text'

        if inputtype == 'textarea':
            html_content += f'''
                    <div class="form-group">
                        <label for="{word}">{new_word.capitalize()}</label>
                        <textarea class="textarea" id="{word}" name="{word}"></textarea>
                    </div>
            '''
        elif inputtype == 'file':
            html_content += f'''
                    <div class="form-group">
                        <label for="{word}">{new_word.capitalize()}</label>
                        <input type="file" id="{word}" name="{word}" accept="image/*">
                    </div>
            '''
        else:
            html_content += f'''
                    <div class="form-group">
                        <label for="{word}">{new_word.capitalize()}</label>
                        <input type="{inputtype}" id="{word}" name="{word}">
                    </div>
            '''

    html_content += '''
                <input type="submit" value="Submit and Download">
            </form>
        </div>
    </body>

    <script>
        $(document).on("keydown", "form", function(event) { 
            return event.key != "Enter";
        });
    </script>

    </html>
    '''

    return html_content


@app.route('/download', methods=['GET'])
def download():
    return send_file('Format_Surat_Lamaran.docx', as_attachment=True)

@app.route('/')
def index():
    data = {}

    symbol = r'{{}}'
    data[0] = r'{{text_nama_field}}'
    data[1] = r'{{text_nama_lengkap}}'
    data[2] = r'{{number_nama_field}}'
    data[3] = r'{{number_nomor_telepon}}'
    data[4] = r'{{file_nama_field}}'
    data[5] = r'{{file_foto_profile}}'

    return render_template('upload.html', data=data, symbol=symbol)

@app.route('/upload',  methods=['POST'])
def process():
    session["typeof"] = None 
    path = 'uploads/'

    if not os.path.exists(path):
        os.makedirs(path) 

    if request.form.get("typeof") == 'lainnya':
        file = request.files['file']

        filename = os.path.join(path, file.filename)
        file.save(filename)
    else:
        session["typeof"] = request.form.get("typeof")
        if request.form.get('typeof') == 'cuti':
            filename = os.path.join(path, "formulir-cuti.docx")
        else:
            filename = os.path.join(path, "Surat-Permohonan-Ijin.docx")

    os.remove(r'filename.json')
    with open('filename.json', 'w') as fl:
        info = {"file": ""}
        json.dump(info, fl)

    with open('filename.json') as fn:
        info = json.load(fn)

    if info["file"] == "":
        info["file"] = filename 

        os.remove(r'filename.json')
        with open('filename.json', 'w') as fl:
            json.dump(info, fl)

    doc = Document(fr'{filename}')
    words = set()
    for paragraph in doc.paragraphs:
        for word in extract_words(paragraph.text):
            print(word)
            words.add(word)
    
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for word in extract_words(paragraph.text):
                        words.add(word)
    
    return generate_html_form(words= list(words))
        
@app.route('/generate-form-file',  methods=['POST'])
def process_api():
    from docx import Document
    
    if request.method == 'POST':
        if request.form['key'] == 'vF6DfT5u0VeA8WEZv7RlMDwumlIHOK':
            file = request.files.getlist('files')
            filename = ""

            for f in file:
                filename = request.form['filename']
                f.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            
            name = filename
            filename = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            doc = Document(fr'{filename}')
            words = set()

            for paragraph in doc.paragraphs:
                for word in extract_words(paragraph.text):
                    words.add(word)
            
                        
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for word in extract_words(paragraph.text):
                                words.add(word)
            
            return {"documentfield": list(words), "filename": name, "status": "success"}
        else:
            return {"status": "Credential failed"}
    else:
        return jsonify({"status": "Upload API GET Request Running"})
    

@app.route('/submit', methods=['POST'])
def submit():
    if request.method == 'POST':

        try:
            with open('filename.json') as fn:
                info = json.load(fn)
            filename = info["file"]

            doc = DocxTemplate(filename)
            data = {}

            os.remove(r'filename.json')
            with open('filename.json', 'w') as fl:
                info = {"file": ""}
                json.dump(info, fl)

            for key, value in request.files.items():
                f = request.files[key]
                FILE_NAME = os.path.join("file/",f.filename) 
                filesaved = f.save(FILE_NAME)
                data = {key: InlineImage(doc, FILE_NAME)}

            for key, value in request.form.items():
                data[key] = value
            print(data)
            name = f"filled_{filename.split('/')[-1]}"
            doc.render(data)
            doc.save(r'filled_'+filename.split('/')[-1])

            # for paragraph in doc.paragraphs:
            #     for run in paragraph.runs:
            #         for k,v in data.items():
            #             print(run.text)
            #             if k in run.text:
            #                 font = run.font
            #                 size = font.size
            #                 run.text = run.text.replace(str(k), str(v))
            #                 run.font.size = size
                        
            # for table in doc.tables:
            #     for row in table.rows:
            #         for cell in row.cells:
            #             for paragraph in cell.paragraphs:
            #                 for run in paragraph.runs:
            #                     for k,v in data.items():
            #                         if k in run.text:
            #                             font = run.font
            #                             size = font.size
            #                             run.text = run.text.replace(str(k), str(v))
            #                             run.font.size = size

            if not session.get("typeof"):
                threading.Thread(target=download_and_redirect1, args=(name,)).start()
            else:
                threading.Thread(target=download_and_redirect2, args=(name,)).start()

            return send_file(name, as_attachment=True)
        except:
            return redirect("/")

@app.route('/submit-api', methods=['POST'])
def submitapi():
    try:
        if request.method == 'POST':
            if request.form['key'] == 'vF6DfT5u0VeA8WEZv7RlMDwumlIHOK':
                filename    = request.form['filename']
                data        = json.loads(request.form['jsondata'])
                folder      = request.form['folder']

                file = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                doc = DocxTemplate(file)
                newdata = {}
                imagedata = {}

                for key, value in data.items():
                    if "file" in key or "image" in key:
                        newval  = value.split('❧')
                        imagedata = {key: InlineImage(doc, os.path.join(folder, newval[0]), width=Mm(int(newval[1])), height=Mm(int(newval[2])))}

                for key, value in data.items():
                    if "file" not in key or "image" not in key:
                        newdata[key] = value
                
                name = f"filled_{filename.split('/')[-1]}"
                document = doc.render(newdata)
                document = doc.render(imagedata)
                doc.save(r'filled_'+filename.split('/')[-1])
                
                documentpdf = Document()
                documentpdf.LoadFromFile(name)
                documentpdf.Watermark = None
                documentpdf.SaveToFile(folder+".pdf", FileFormat.PDF)
                documentpdf.Close()

                threading.Thread(target=download_and_redirect_api, args=(name, folder,)).start()
                return send_file(name, as_attachment=True)
                #return {"jsondata": data, "filename": filename, "status": "success"}
            else:
                return {"status": "Credential failed"}
    except:
        return {"status": "server error"}

@app.route('/upload-image-api',  methods=['POST'])
def uploadimage_api():
    try:
        if request.method == 'POST':
            if request.form['key'] == 'vF6DfT5u0VeA8WEZv7RlMDwumlIHOK':
                file = request.files.getlist('files')
                os.makedirs(request.form['folder'], exist_ok=True) 

                for f in file:
                    filename = request.form['filename']
                    f.save(os.path.join(request.form['folder'], filename))
                
                return {"status": "success"}
            else:
                return {"status": "Credential failed"}
        else:
            return jsonify({"status": "Upload API GET Request Running"})
    except:
        return {"status": "server error"}
    
@app.route('/download-pdf',  methods=['POST'])
def downloadpdf():
    try:
        if request.method == 'POST':
            if request.form['key'] == 'vF6DfT5u0VeA8WEZv7RlMDwumlIHOK':
                input_pdf = request.form['pdfname']
                output_pdf = 'filled_'+request.form['pdfname']
                text_to_remove = "Evaluation Warning: The document was created with Spire.Doc for Python."  # Modify with the specific text you want to remove

                remove_text_from_pdf(input_pdf, output_pdf, text_to_remove)

                threading.Thread(target=download_and_redirect_pdf, args=(input_pdf, output_pdf,)).start()
                return send_file(output_pdf, as_attachment=True)
            else:
                return {"status": "Credential failed"}
        else:
            return jsonify({"status": "Upload API GET Request Running"})
    except:
        return {"status": "server error"}

def remove_text_from_pdf(input_pdf, output_pdf, text_to_remove):
    doc = fitz.open(input_pdf)
    
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        
        blocks = page.get_text("dict")["blocks"]
        
        for block in blocks:
            if block["type"] == 0:  # 0 represents text block type
                for line in block["lines"]:
                    for span in line["spans"]:
                        if text_to_remove in span["text"]:
                            rect = fitz.Rect(span["bbox"])
                            page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
    doc.save(output_pdf)

if __name__ == '__main__':
    app.run(debug=True)
    uvicorn.run("main:app", host = "192.1.2.3", port = 5050, log_level = "info", reload = True)